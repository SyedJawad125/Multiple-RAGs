"""
Resume Parser Agent
───────────────────
1. ResumeFileExtractor   — pulls raw text from PDF / DOCX / DOC
2. ResumeParserAgent     — sends raw text to GPT-4o, gets structured JSON back
"""
import json
import logging
import time
from django.conf import settings
from tenacity import retry, stop_after_attempt, wait_exponential

logger = logging.getLogger(__name__)

# ── Prompt ─────────────────────────────────────────────────────────────────────
PARSE_SYSTEM = (
    'You are an expert resume parser. '
    'Extract every piece of information accurately. '
    'Always respond with valid JSON only — no markdown, no explanation.'
)

PARSE_PROMPT = """\
Parse the resume below and return a JSON object with EXACTLY this structure:

{{
  "candidate_name":     "Full Name",
  "candidate_email":    "email@example.com",
  "candidate_phone":    "+1-555-000-0000",
  "candidate_location": "City, Country",
  "candidate_linkedin": "https://linkedin.com/in/...",
  "candidate_github":   "https://github.com/...",
  "candidate_website":  "https://...",

  "total_experience_years": 5.0,
  "highest_education": "bachelor",

  "education_details": [
    {{
      "degree":      "Bachelor of Science",
      "field":       "Computer Science",
      "institution": "MIT",
      "year":        2020,
      "gpa":         "3.9"
    }}
  ],

  "experience_details": [
    {{
      "title":        "Senior Backend Engineer",
      "company":      "Acme Corp",
      "start_date":   "2021-01",
      "end_date":     "2024-03",
      "years":        3.2,
      "is_current":   false,
      "description":  "Led a team of 5, built microservices…",
      "technologies": ["Python", "Django", "PostgreSQL", "AWS"]
    }}
  ],

  "skills": [
    {{
      "name":        "Python",
      "category":    "technical",
      "proficiency": "expert",
      "years_used":  6.0
    }}
  ],

  "certifications": [
    {{
      "name":   "AWS Solutions Architect Associate",
      "issuer": "Amazon",
      "year":   2022
    }}
  ],

  "languages": ["English", "French"],
  "summary":   "Brief professional summary…"
}}

highest_education must be one of: high_school | associate | bachelor | master | mba | phd | other

Resume text:
{resume_text}
"""


# ── File extractor ─────────────────────────────────────────────────────────────

class ResumeFileExtractor:
    """Extracts raw text from PDF, DOCX, DOC files."""

    def extract(self, file_path: str, file_type: str) -> str:
        ext = file_type.lower().lstrip('.')
        if ext == 'pdf':
            return self._pdf(file_path)
        if ext in ('docx', 'doc'):
            return self._docx(file_path)
        raise ValueError(f'Unsupported file type: {file_type}')

    def _pdf(self, path: str) -> str:
        """pdfplumber first, PyMuPDF as fallback."""
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(path) as pdf:
                for p in pdf.pages:
                    t = p.extract_text()
                    if t:
                        pages.append(t)
            text = '\n'.join(pages)
            if text.strip():
                return text
        except Exception as e:
            logger.debug(f'pdfplumber failed ({e}), trying PyMuPDF')

        import fitz  # pymupdf
        doc = fitz.open(path)
        text = '\n'.join(p.get_text() for p in doc)
        doc.close()
        return text

    def _docx(self, path: str) -> str:
        from docx import Document
        doc    = Document(path)
        parts  = [p.text for p in doc.paragraphs if p.text.strip()]
        # also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return '\n'.join(parts)


# ── Parser agent ───────────────────────────────────────────────────────────────

class ResumeParserAgent:
    """Sends raw resume text to the LLM and returns structured data."""

    def __init__(self):
        from openai import OpenAI
        self._client = OpenAI(api_key=settings.OPENAI_API_KEY)
        self._model  = settings.OPENAI_MODEL

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(min=2, max=10))
    def parse(self, raw_text: str, resume_id: str) -> dict:
        """
        Returns:
            {'success': True,  'data': {...}, 'tokens_used': N, 'ms': N}
            {'success': False, 'error': '...', 'tokens_used': N}
        """
        t0 = time.time()
        tokens_used = 0
        try:
            resp = self._client.chat.completions.create(
                model           = self._model,
                messages        = [
                    {'role': 'system', 'content': PARSE_SYSTEM},
                    {'role': 'user',   'content': PARSE_PROMPT.format(resume_text=raw_text[:10000])},
                ],
                temperature     = 0.0,
                max_tokens      = 3000,
                response_format = {'type': 'json_object'},
            )
            tokens_used = resp.usage.total_tokens
            data        = json.loads(resp.choices[0].message.content)
            ms          = int((time.time() - t0) * 1000)
            logger.info(f'[ResumeParser] {resume_id} parsed — {tokens_used} tokens, {ms} ms')
            self._log(resume_id, 'success', data, tokens_used, ms)
            return {'success': True, 'data': data, 'tokens_used': tokens_used, 'ms': ms}

        except json.JSONDecodeError as e:
            ms = int((time.time() - t0) * 1000)
            logger.error(f'[ResumeParser] JSON decode error for {resume_id}: {e}')
            self._log(resume_id, 'error', {}, tokens_used, ms, str(e))
            return {'success': False, 'error': str(e), 'tokens_used': tokens_used}

        except Exception as e:
            ms = int((time.time() - t0) * 1000)
            logger.error(f'[ResumeParser] Error for {resume_id}: {e}')
            self._log(resume_id, 'error', {}, tokens_used, ms, str(e))
            raise  # let tenacity retry

    def _log(self, resume_id, status, output, tokens, ms, error=''):
        try:
            from apps.screening.models import AgentExecutionLog
            AgentExecutionLog.objects.create(
                agent_type         = 'resume_parser',
                input_summary      = {'resume_id': resume_id},
                output_summary     = {k: output[k] for k in ['candidate_name', 'candidate_email',
                                                               'total_experience_years', 'highest_education']
                                      if k in output},
                status             = status,
                error_message      = error,
                tokens_used        = tokens,
                processing_time_ms = ms,
                model_used         = self._model,
            )
        except Exception as e:
            logger.warning(f'Failed to write AgentExecutionLog: {e}')