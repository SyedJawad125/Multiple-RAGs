# apps/ragstack/document_processor.py - FIXED VERSION
"""
===============================================================
Document Processor - Extract, Chunk, Embed - FIXED
===============================================================
Processes uploaded documents WITHOUT async issues
===============================================================
"""

import logging
import hashlib
import io
from typing import List
from django.core.files.storage import default_storage
from django.db import transaction
from django.utils import timezone
import PyPDF2
import docx
import csv
import openpyxl

from .models import Document, DocumentChunk
from .rag_service import EmbeddingService

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various document formats"""
    
    @staticmethod
    def extract_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
                        logger.info(f"Extracted text from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"✅ Extracted {len(full_text)} characters from PDF")
            return full_text.strip()
        
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    paragraphs.append(row_text)
            
            full_text = "\n\n".join(paragraphs)
            logger.info(f"✅ Extracted {len(full_text)} characters from DOCX")
            return full_text.strip()
        
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_from_txt(file_content: bytes) -> str:
        """Extract text from TXT"""
        try:
            text = file_content.decode('utf-8')
            logger.info(f"✅ Extracted {len(text)} characters from TXT")
            return text.strip()
        except UnicodeDecodeError:
            # Try alternative encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = file_content.decode(encoding)
                    logger.info(f"✅ Extracted {len(text)} characters from TXT ({encoding})")
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Failed to decode text file with any common encoding")
    
    @staticmethod
    def extract_from_csv(file_content: bytes) -> str:
        """Extract text from CSV"""
        try:
            content = file_content.decode('utf-8')
            csv_reader = csv.reader(io.StringIO(content))
            
            rows = []
            for row in csv_reader:
                rows.append(" | ".join(row))
            
            full_text = "\n".join(rows)
            logger.info(f"✅ Extracted {len(full_text)} characters from CSV")
            return full_text
        
        except Exception as e:
            logger.error(f"CSV extraction failed: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to extract text from CSV: {str(e)}")
    
    @staticmethod
    def extract_from_xlsx(file_content: bytes) -> str:
        """Extract text from XLSX"""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
            
            all_text = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                all_text.append(f"Sheet: {sheet_name}\n")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        all_text.append(row_text)
                
                all_text.append("\n")
            
            full_text = "\n".join(all_text)
            logger.info(f"✅ Extracted {len(full_text)} characters from XLSX")
            return full_text
        
        except Exception as e:
            logger.error(f"XLSX extraction failed: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to extract text from XLSX: {str(e)}")


class TextChunker:
    """Intelligent text chunking with overlap"""
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        separators: List[str] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", ". ", " "]
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap"""
        if not text or len(text) == 0:
            logger.warning("Empty text provided for chunking")
            return []
        
        logger.info(f"Chunking text of length {len(text)}")
        
        chunks = []
        current_chunk = ""
        
        # Split by separators
        parts = self._split_by_separators(text)
        logger.info(f"Split into {len(parts)} parts")
        
        for part in parts:
            # If part itself is too long, split it
            if len(part) > self.chunk_size:
                sub_chunks = self._split_long_part(part)
                chunks.extend(sub_chunks)
                current_chunk = ""
                continue
            
            # Add part to current chunk
            if len(current_chunk) + len(part) <= self.chunk_size:
                current_chunk += part
            else:
                # Save current chunk and start new one with overlap
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Get overlap from end of current chunk
                overlap = self._get_overlap(current_chunk)
                current_chunk = overlap + part
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        logger.info(f"✅ Created {len(chunks)} chunks")
        return chunks
    
    def _split_by_separators(self, text: str) -> List[str]:
        """Split text using separator hierarchy"""
        parts = [text]
        
        for separator in self.separators:
            new_parts = []
            for part in parts:
                if separator in part:
                    split_parts = part.split(separator)
                    for i, sp in enumerate(split_parts):
                        if i < len(split_parts) - 1:
                            new_parts.append(sp + separator)
                        else:
                            new_parts.append(sp)
                else:
                    new_parts.append(part)
            parts = new_parts
        
        return [p for p in parts if p.strip()]
    
    def _split_long_part(self, part: str) -> List[str]:
        """Split a part that's longer than chunk_size"""
        chunks = []
        start = 0
        
        while start < len(part):
            end = min(start + self.chunk_size, len(part))
            chunk = part[start:end]
            chunks.append(chunk.strip())
            start = end - self.chunk_overlap
        
        return chunks
    
    def _get_overlap(self, text: str) -> str:
        """Get overlap text from end of chunk"""
        if len(text) <= self.chunk_overlap:
            return text
        return text[-self.chunk_overlap:]


class DocumentProcessor:
    """Main document processor - SYNCHRONOUS VERSION"""
    
    def __init__(self):
        self.text_extractor = TextExtractor()
        self.text_chunker = TextChunker()
        self.embedding_service = EmbeddingService()
    
    def process_document(self, document_id: str):
        """
        Process document: extract, chunk, embed, store - SYNC VERSION
        """
        logger.info(f"🚀 Starting to process document {document_id}")
        
        try:
            # Get document - Use .get() not .aget()
            document = Document.objects.get(id=document_id)
            document.status = 'processing'
            document.save(update_fields=['status'])
            
            logger.info(f"📄 Processing: {document.filename}")
            
            # Step 1: Read file
            with default_storage.open(document.file_path, 'rb') as f:
                file_content = f.read()
            
            logger.info(f"📖 Read {len(file_content)} bytes")
            
            # Step 2: Extract text
            logger.info(f"🔍 Extracting text from {document.filename}")
            text = self._extract_text(file_content, document.content_type)
            
            if not text or len(text.strip()) < 10:
                raise ValueError("Extracted text is too short or empty")
            
            logger.info(f"✅ Extracted {len(text)} characters")
            
            # Step 3: Chunk text
            logger.info(f"✂️  Chunking text...")
            chunks = self.text_chunker.chunk_text(text)
            
            if not chunks:
                raise ValueError("No chunks generated from text")
            
            logger.info(f"✅ Generated {len(chunks)} chunks")
            
            # Step 4: Generate embeddings
            logger.info("🧠 Generating embeddings...")
            embeddings = self.embedding_service.embed_texts(chunks)
            logger.info(f"✅ Generated {len(embeddings)} embeddings")
            
            # Step 5: Store chunks with embeddings
            logger.info("💾 Storing chunks in database...")
            self._store_chunks(document, chunks, embeddings)
            
            # Step 6: Update document status
            document.status = 'completed'
            document.chunks_count = len(chunks)
            document.processed_at = timezone.now()
            document.metadata.update({
                'text_length': len(text),
                'chunks_count': len(chunks),
                'processing_completed': True
            })
            document.save(update_fields=['status', 'chunks_count', 'processed_at', 'metadata'])
            
            logger.info(f"✅ Document {document_id} processed successfully! ({len(chunks)} chunks)")
        
        except Exception as e:
            logger.error(f"❌ Document processing failed: {str(e)}", exc_info=True)
            
            # Update document status to failed
            try:
                document = Document.objects.get(id=document_id)
                document.status = 'failed'
                if not document.metadata:
                    document.metadata = {}
                document.metadata['error'] = str(e)
                document.save(update_fields=['status', 'metadata'])
            except Exception:
                pass
    
    def _extract_text(self, file_content: bytes, content_type: str) -> str:
        """Extract text based on content type"""
        logger.info(f"Extracting text for content type: {content_type}")
        
        if content_type == 'application/pdf':
            return self.text_extractor.extract_from_pdf(file_content)
        elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self.text_extractor.extract_from_docx(file_content)
        elif content_type == 'text/plain':
            return self.text_extractor.extract_from_txt(file_content)
        elif content_type == 'text/csv':
            return self.text_extractor.extract_from_csv(file_content)
        elif content_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            return self.text_extractor.extract_from_xlsx(file_content)
        else:
            raise ValueError(f"Unsupported content type: {content_type}")
    
    def _store_chunks(
        self,
        document: Document,
        chunks: List[str],
        embeddings: List[List[float]]
    ):
        """Store chunks with embeddings in database"""
        chunk_objects = []
        
        logger.info(f"Preparing {len(chunks)} chunks for storage...")
        
        for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
            # Generate content hash
            content_hash = hashlib.md5(chunk_text.encode()).hexdigest()
            
            chunk_obj = DocumentChunk(
                document=document,
                content=chunk_text,
                chunk_index=idx,
                embedding=embedding,
                content_hash=content_hash,
                word_count=len(chunk_text.split()),
                metadata={
                    'chunk_length': len(chunk_text),
                    'word_count': len(chunk_text.split())
                }
            )
            chunk_objects.append(chunk_obj)
        
        # Bulk create - SYNC version
        DocumentChunk.objects.bulk_create(chunk_objects, batch_size=100)
        
        logger.info(f"✅ Stored {len(chunk_objects)} chunks for document {document.id}")


# ============================================================
# STANDALONE TEST FUNCTION
# ============================================================

def test_document_processor():
    """Test the document processor"""
    import sys
    import os
    import django
    
    # Setup Django
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    sys.path.insert(0, project_root)
    os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
    django.setup()
    
    from apps.ragstack.models import Document
    
    # Get a document to process
    docs = Document.objects.filter(status='uploading').order_by('-created_at')
    
    if not docs.exists():
        print("❌ No documents to process")
        return
    
    doc = docs.first()
    print(f"🔄 Processing document: {doc.filename}")
    
    processor = DocumentProcessor()
    processor.process_document(str(doc.id))
    
    print("✅ Done!")


if __name__ == "__main__":
    test_document_processor()